"""Generate a Word document technical report for CHIB system."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Title Page ──
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('CHAK Health Information Bridge\n(CHIB)')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Technical Overview & System Architecture')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph('')
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run(f'Prepared: {datetime.date.today().strftime("%B %d, %Y")}')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph('')
audience = doc.add_paragraph()
audience.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = audience.add_run('For: Executive Leadership — CHAK')
run.font.size = Pt(12)
run.italic = True

doc.add_page_break()

# ── Table of Contents (manual) ──
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. What Is CHIB?',
    '2. System Architecture (How It All Connects)',
    '3. Where the Data Comes From',
    '4. Technologies Used',
    '5. The Two Applications',
    '6. Key Features',
    '7. How DHIS2 Integration Works',
    '8. Deployment & Hosting',
    '9. Security',
    '10. Summary',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 1. WHAT IS CHIB
# ══════════════════════════════════════════════════════════════════
doc.add_heading('1. What Is CHIB?', level=1)
doc.add_paragraph(
    'CHIB (CHAK Health Information Bridge) is a web-based system that connects to Kenya\'s '
    'national health database (DHIS2) to display real-time HIV/AIDS program data. '
    'It was built for the Christian Health Association of Kenya (CHAK) to help '
    'health managers track key indicators like:'
)
bullets = [
    'How many patients are newly starting HIV treatment (TX_NEW)',
    'How many patients are currently on treatment (TX_CURR)',
    'Viral load testing and suppression rates',
    'HIV testing numbers across different entry points',
    'PrEP (prevention) programme statistics',
    'TB/HIV co-infection management',
    'Cervical cancer screening among HIV-positive women',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph(
    'In simple terms: CHIB takes data from the national DHIS2 server and presents '
    'it in easy-to-read charts and dashboards — accessible from any web browser.'
)

# ══════════════════════════════════════════════════════════════════
# 2. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════
doc.add_heading('2. System Architecture (How It All Connects)', level=1)
doc.add_paragraph(
    'The system follows a simple three-layer architecture:'
)

# Browser layer
doc.add_heading('Layer 1: Web Browser (Frontend)', level=2)
doc.add_paragraph(
    'The user opens a website in their browser. The site is built with ordinary '
    'JavaScript (no heavy frameworks) and uses Chart.js and Highcharts to draw '
    'beautiful charts. The interface is styled with Tailwind CSS for a modern, '
    'clean look.'
)

# Server layer
doc.add_heading('Layer 2: Application Server (Backend)', level=2)
doc.add_paragraph(
    'The server is written in Python using Flask (a lightweight web framework). '
    'It handles user requests, talks to the DHIS2 server to fetch data, processes '
    'that data, and sends it back to the browser as JSON. It also has an AI '
    'assistant powered by Google Gemini.'
)

# DHIS2 layer
doc.add_heading('Layer 3: DHIS2 Server (Data Source)', level=2)
doc.add_paragraph(
    'The national DHIS2 server is hosted at ereporting.chak.or.ke. It stores all '
    'the health facility data. CHIB sends queries to this server asking for specific '
    'indicators, and the DHIS2 server sends back the numbers.'
)

doc.add_paragraph('')
doc.add_heading('Visual Diagram', level=2)
doc.add_paragraph(
    'User\'s Browser  →  CHIB Server (Python/Flask on Render.com)  →  DHIS2 Server (ereporting.chak.or.ke)  →  Response with Data  →  Charts Displayed'
)

p = doc.add_paragraph()
run = p.add_run('The AI assistant also uses Google Gemini:')
run.italic = True
doc.add_paragraph(
    'User Question  →  CHIB Server  →  Google Gemini AI  →  SQL Query  →  Local Database  →  Answer'
)

# ══════════════════════════════════════════════════════════════════
# 3. WHERE THE DATA COMES FROM
# ══════════════════════════════════════════════════════════════════
doc.add_heading('3. Where the Data Comes From', level=1)
doc.add_paragraph(
    'All health data is fetched live from DHIS2 — Kenya\'s national health information '
    'system. There is no copying or batch processing. When you open a dashboard, '
    'the system goes directly to DHIS2 and asks for the latest numbers.'
)

doc.add_paragraph('The system covers four counties:')
counties = [
    ('Meru', 'The primary focus area with the most facilities'),
    ('Embu', 'Secondary coverage area'),
    ('Nyandarua', 'Tertiary coverage area'),
    ('Tharaka Nithi', 'Quaternary coverage area'),
]
for name, desc in counties:
    doc.add_paragraph(f'{name} — {desc}', style='List Bullet')

doc.add_paragraph(
    'Users can drill down from county level → to sub-county level → to individual '
    'health facilities to see data at every level.'
)

# ══════════════════════════════════════════════════════════════════
# 4. TECHNOLOGIES USED
# ══════════════════════════════════════════════════════════════════
doc.add_heading('4. Technologies Used', level=1)
doc.add_paragraph('Here is a simple breakdown of every technology used:')

# Table
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Technology'
hdr_cells[1].text = 'What It Is'
hdr_cells[2].text = 'What It Does Here'

techs = [
    ('Python', 'Programming language', 'The brain of the system — all logic is written in Python'),
    ('Flask', 'Web framework', 'Handles all web requests and serves pages to the browser'),
    ('JavaScript', 'Programming language', 'Runs in the browser to make charts interactive'),
    ('Chart.js', 'Charting library', 'Draws bar charts, doughnut charts, and population pyramids'),
    ('Highcharts', 'Charting library', 'Draws line charts for trends over time'),
    ('Tailwind CSS', 'Styling framework', 'Makes the website look clean and professional'),
    ('Gunicorn', 'Web server', 'Runs the Flask app in production (handles multiple users)'),
    ('Google Gemini AI', 'Artificial intelligence', 'Powers the AI assistant that answers questions in plain English'),
    ('Pandas', 'Data analysis library', 'Organizes and processes data from DHIS2 into useful formats'),
    ('SQLite', 'Database', 'Stores clinic expenditure data locally for AI queries'),
    ('Render.com', 'Cloud hosting', 'Hosts the live website on the internet (free plan)'),
]
for tech, what, does in techs:
    row_cells = table.add_row().cells
    row_cells[0].text = tech
    row_cells[1].text = what
    row_cells[2].text = does

# ══════════════════════════════════════════════════════════════════
# 5. THE TWO APPLICATIONS
# ══════════════════════════════════════════════════════════════════
doc.add_heading('5. The Two Applications', level=1)
doc.add_paragraph('The system is split into two separate websites:')

doc.add_heading('App 1: DHIS2 Superpower (Streamlit)', level=2)
doc.add_paragraph('URL: https://dhis2-superpower.onrender.com')
bullets = [
    'Built with Streamlit (a Python tool for making simple web apps)',
    'Allows users to type questions in plain English, e.g., "Show me HTS_TST_POS for Meru"',
    'Uses Google Gemini AI to translate English into DHIS2 queries',
    'Can export results to CSV',
    'Best for: Ad-hoc data requests and exploration',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('App 2: CHIB Train Dashboard (Flask)', level=2)
doc.add_paragraph('URL: https://dhis2-superpower-1.onrender.com')
bullets = [
    'Built with Flask (a more powerful Python web framework)',
    'Full executive dashboard with interactive charts and filters',
    'Multiple tabs for HIV Treatment, HIV Testing, PrEP, TB, and more',
    'Real-time DHIS2 data connection',
    'AI Executive Copilot for natural-language data queries',
    'PBIX dashboard integration for legacy Power BI content',
    'Best for: Routine monitoring, executive reporting, and daily use',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# ══════════════════════════════════════════════════════════════════
# 6. KEY FEATURES
# ══════════════════════════════════════════════════════════════════
doc.add_heading('6. Key Features', level=1)

features = [
    ('Live DHIS2 Data', 'Every time you open a page, it fetches fresh data directly from the national DHIS2 server. No stale data, no manual exports.'),
    ('Interactive Charts', 'Data is presented as bar charts, line charts, doughnut charts, and population pyramids. Hover over any chart to see exact numbers.'),
    ('Multi-Level Drill-Down', 'Start at the national/county level, then drill into sub-county, then into individual health facilities.'),
    ('AI Executive Copilot', 'Click the "AI Assist" button and type questions like "Which 3 clinics have the highest expenditure?" The AI understands your question, looks up the data, and gives you an answer.'),
    ('Multiple Program Areas', 'The dashboard covers HIV Treatment, HIV Testing Services (HTS), PrEP prevention, TB/HIV co-management, Cervical Cancer screening, and more.'),
    ('PBIX Dashboard Compatibility', 'Legacy Power BI dashboards are still accessible through the system, ensuring no loss of historical reporting.'),
    ('CSV Export', 'Any data visible on screen can be exported to CSV for further analysis in Excel.'),
    ('Mobile-Friendly', 'The interface works on phones and tablets, not just desktop computers.'),
]
for title_text, desc in features:
    doc.add_heading(title_text, level=2)
    doc.add_paragraph(desc)

# ══════════════════════════════════════════════════════════════════
# 7. HOW DHIS2 INTEGRATION WORKS
# ══════════════════════════════════════════════════════════════════
doc.add_heading('7. How DHIS2 Integration Works', level=1)
doc.add_paragraph(
    'DHIS2 (District Health Information System 2) is Kenya\'s national health data '
    'system used by all health facilities to report their monthly numbers.'
)

doc.add_heading('Step-by-Step Process', level=2)
steps = [
    'A user selects filters on the dashboard (e.g., "Show TX_NEW for Meru County for the last 12 months")',
    'The system looks up the DHIS2 codes for the selected county and indicator',
    'It sends a request to DHIS2\'s analytics API at http://ereporting.chak.or.ke:8500/api',
    'DHIS2 processes the request and returns the numbers',
    'The system organizes the raw numbers into a useful format (trends, age groups, gender splits)',
    'The browser displays the data as beautiful, interactive charts',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_paragraph('')
doc.add_heading('What Indicators Are Tracked', level=2)
doc.add_paragraph('The system monitors over 50 different DHIS2 data elements. The most important ones are:')

indicators = [
    ('TX_NEW', 'Patients newly enrolled on ART (antiretroviral therapy)'),
    ('TX_CURR', 'Patients currently receiving ART'),
    ('TX_PVLS', 'Viral load testing — checking if treatment is working'),
    ('HTS_TST', 'HIV testing services — how many people were tested'),
    ('HTS_POSITIVE', 'People who tested HIV positive'),
    ('PrEP_NEW', 'People starting PrEP (prevention medication)'),
    ('PrEP_CURR', 'People currently on PrEP'),
    ('TB_STAT', 'TB screening among HIV patients'),
    ('CACX Screening', 'Cervical cancer screening for HIV-positive women'),
    ('TX_ML (IIT)', 'Patients who missed appointments (interruption in treatment)'),
]
for name, desc in indicators:
    doc.add_paragraph(f'{name} — {desc}', style='List Bullet')

# ══════════════════════════════════════════════════════════════════
# 8. DEPLOYMENT & HOSTING
# ══════════════════════════════════════════════════════════════════
doc.add_heading('8. Deployment & Hosting', level=1)
doc.add_paragraph(
    'Both applications are hosted on Render.com, a cloud platform similar to Heroku. '
    'The system runs on Render\'s free tier.'
)

doc.add_heading('Hosting Specifications', level=2)
specs = [
    ('Platform', 'Render.com (Free Tier)'),
    ('Memory', '512 MB RAM'),
    ('Processor', '0.1 vCPU (shared)'),
    ('Storage', 'Ephemeral — data resets on each deploy'),
    ('Auto-Scale', 'No — single instance per service'),
    ('SSL/HTTPS', 'Automatic — managed by Render'),
    ('Source Control', 'GitHub (main branch) — every push triggers a deploy'),
]
for label, value in specs:
    p = doc.add_paragraph()
    run = p.add_run(f'{label}: ')
    run.bold = True
    p.add_run(value)

doc.add_paragraph('')
doc.add_heading('Limitations of Free Tier', level=2)
bullets = [
    'The app "spins down" (goes to sleep) after about 15 minutes of inactivity',
    'First visit after inactivity takes 30–50 seconds to wake up (cold start)',
    'Limited to 512 MB RAM — may be slow with very large data requests',
    'No custom domain — uses render.com subdomain',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')
doc.add_paragraph(
    'An upgrade to the Starter plan ($7/month) would eliminate the spin-down delay '
    'and provide faster performance.'
)

# ══════════════════════════════════════════════════════════════════
# 9. SECURITY
# ══════════════════════════════════════════════════════════════════
doc.add_heading('9. Security', level=1)
doc.add_paragraph('The system uses the following security measures:')

bullets = [
    'DHIS2 credentials are stored as environment variables on Render — never written in code or committed to GitHub',
    'All web traffic uses HTTPS (encrypted) between the user and the server',
    'The AI feature uses API keys stored securely in environment variables',
    'Multiple Gemini API keys are configured for redundancy (if one fails, the next is used)',
    'No user data is stored — all health data stays in DHIS2 and is queried live',
    'The system uses the same authentication as the national DHIS2 server (username & password)',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# ══════════════════════════════════════════════════════════════════
# 10. SUMMARY
# ══════════════════════════════════════════════════════════════════
doc.add_heading('10. Summary', level=1)
doc.add_paragraph(
    'CHIB (CHAK Health Information Bridge) is a modern, web-based dashboard that '
    'brings DHIS2 health data to life through interactive charts and AI-powered '
    'analytics. It connects directly to Kenya\'s national DHIS2 server, fetches '
    'real-time data, and presents it in an easy-to-understand format.'
)
doc.add_paragraph(
    'Built with Python (Flask), JavaScript (Chart.js), and Google Gemini AI, the '
    'system covers four counties and tracks key HIV/AIDS program indicators including '
    'treatment, testing, prevention, TB co-infection, and cervical cancer screening.'
)
doc.add_paragraph(
    'Both applications are hosted on Render.com\'s free tier and are accessible '
    'from any web browser. The system is designed to be simple, maintainable, '
    'and easy for non-technical users to navigate.'
)

doc.add_paragraph('')
doc.add_paragraph('')
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('— End of Report —')
run.italic = True
run.font.color.rgb = RGBColor(150, 150, 150)

# ── Save ──
output_path = 'c:\\Users\\ADMIN\\Documents\\dhis2_superpower\\CHIB_Technical_Report.docx'
doc.save(output_path)
print(f'✅ Report saved to: {output_path}')
